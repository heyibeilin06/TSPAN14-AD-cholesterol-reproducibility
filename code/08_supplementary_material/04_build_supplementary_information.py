#!/usr/bin/env python3
"""Author Supplementary Information v19 from a blank Word document."""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "outputs" / "supplement_v19"
TITLE = "A non-APOE Alzheimer disease-cholesterol locus converges on TSPAN14 splice choice"

METHODS = [
    ("Study design and evidence organization", [
        "The supplementary analyses were organized around a prespecified evidence sequence: genome-wide Alzheimer disease (AD)-lipid sharing, sensitivity to the extended apolipoprotein E (APOE) region, non-APOE regional prioritization, resolution of the TSPAN14 molecular phenotype, exact-event cross-tissue consistency, local causal-scope analyses, cell-context interpretation and transcript-to-protein annotation. Genetic association, molecular quantitative-trait-locus (QTL), disease-state RNA and perturbation evidence were analysed as distinct evidence classes and were combined only after allele, genome-build and molecular-event harmonization.",
        "The exact molecular phenotype was defined by its GRCh38 donor-acceptor coordinates (chr10:80509471-80512144), corresponding to the canonical TSPAN14 exon5-6 junction and amino-acid transition 150/151 in transcript ENST00000429989 (TSPAN14-207). Gene-level or adjacent-junction observations were not treated as exact replication of this event."
    ]),
    ("GWAS resources, harmonization and baseline genetic correlation", [
        "Summary statistics for AD and five lipid traits were obtained from the public resources listed in Supplementary Table S1. Variants were harmonized by chromosome, position and allele, with strand-ambiguous variants removed when orientation could not be resolved. Effect estimates were aligned to the reported effect allele for genome-wide analyses and subsequently reoriented to the inferred AD-risk-increasing allele for locus-level direction comparisons.",
        "Pairwise linkage disequilibrium score regression was applied to estimate genome-wide genetic correlation between AD and high-density lipoprotein cholesterol, low-density lipoprotein cholesterol, triglycerides, total cholesterol and non-high-density lipoprotein cholesterol. Regression intercepts and standard errors were retained from the complete genome-wide models; two-sided P values were used without reclassifying non-significant lipid pairs as evidence of no biological relationship."
    ]),
    ("Extended-APOE conditional and physical-window sensitivity analyses", [
        "The APOE analysis was treated as a sensitivity analysis rather than proof that APOE effects had been completely removed. The extended 19q13 linkage-disequilibrium block was conditioned on the strongest available trait-specific regional signals and on the union of AD and lipid lead signals using the same ancestry-matched linkage-disequilibrium reference used for summary-statistic conditioning. Residualized summary statistics were then re-entered into the genetic-correlation workflow.",
        "A complementary physical-window analysis removed variants in progressively broader regions around the APOE anchor, extending to 5 Mb. Concordance between LD-conditioned estimates and the extreme 5-Mb exclusion was used to evaluate whether the AD-HDL-C correlation was entirely attributable to the dominant APOE neighbourhood. The physical-window result was not interpreted as a formal conditional estimate."
    ]),
    ("Non-APOE regional screening, colocalization and fine-mapping", [
        "Outside chromosome 19, regional AD-lipid pairs were screened using posterior probabilities for distinct and shared association models. Candidate regions were ranked by the probability of a shared signal and by reproducibility across cholesterol traits. Multiple-signal colocalization was evaluated with SuSiE-based models using a common local variant set and ancestry-matched linkage disequilibrium. HyPrColoc was used as a complementary multi-trait clustering analysis; its regional probability was interpreted as evidence for a shared association configuration, not as proof of a serial biological pathway.",
        "Trait-specific credible sets and posterior inclusion probabilities were compared across AD, cholesterol traits and the exact splice phenotype. For approximate-Bayes-factor colocalization, the default priors were p1 = p2 = 1 × 10⁻⁴ and p12 = 1 × 10⁻⁵; p12 was then decreased and increased tenfold while variant Bayes factors were held fixed. This sensitivity analysis covered both regional AD–cholesterol comparisons and the coordinate-identical canonical exon5–6 event paired with AD and each cholesterol trait. The number of allele-compatible overlapping variants was reported for every trait pair. Because lead variants differed by trait and no variant dominated every model, the locus was interpreted at the regulatory-haplotype level. Functional annotation was used to identify enhancer-related candidates and experimentally interrogated variants, but did not convert any single nucleotide into an assumed causal variant."
    ]),
    ("Canonical–cryptic acceptor count modelling", [
        "The local event was represented by canonical and competing cryptic-acceptor reads for 147 donors. Genotypes were coded as 0, 1 or 2 copies of the rs7080009 AD-risk C allele on the GRCh38 forward strand. The C allele is the reverse-complement equivalent of the risk-aligned G allele used in manuscript-facing harmonization. Canonical read fraction was retained as a descriptive measure of local acceptor balance rather than as a full-transcript percent-spliced-in estimate.",
        "A beta-binomial model related cryptic reads to total canonical-plus-cryptic reads and accommodated extra-binomial dispersion. Because no cryptic reads were detected among the 106 C/C donors, donor-level detection was additionally analysed with Firth penalized logistic regression. Profile-likelihood confidence intervals and penalized likelihood-ratio P values were reported. A second Firth model adjusted for standardized log local-cluster depth; progressively stricter depth thresholds and leave-one-donor-out analyses tested robustness to sparse competing-junction counts."
    ]),
    ("Exact-event cross-tissue consistency and adjacent-junction co-usage", [
        "Cross-tissue matching required the same GRCh38 donor and acceptor coordinates as the primary exon5-6 event. Associations were harmonized to the AD-risk-increasing allele and evaluated across anterior cingulate cortex, hippocampus, putamen and cervical spinal cord. Because these GTEx tissues shared donors, they were regarded as repeated biological contexts for a correlated haplotype rather than independent genetic replications.",
        "The previously identified exon6-7 feature was evaluated separately. Across neural samples with available junction counts, Spearman correlation was calculated between log-transformed exon5-6 and exon6-7 counts within each tissue and across the pooled dataset. Positive co-usage argues against a simple mutually exclusive relationship at the aggregate-count level, but does not establish that both junctions occur in the same RNA molecule."
    ]),
    ("Exact-event colocalization, cis-MR and overlap sensitivity", [
        "Full cis-association statistics for the exact exon5-6 event were aligned with AD and cholesterol GWAS variants before single- and multiple-signal colocalization. Linkage-disequilibrium-aware generalized inverse-variance weighted cis-Mendelian randomization (cis-MR) used five genome-wide significant correlated sQTL variants and a signed European-ancestry linkage-disequilibrium matrix rebuilt from phased 1000 Genomes Project data after allele remapping. A lead-variant Wald estimate was retained as a one-instrument sensitivity analysis. Matrix eigenvalues, condition number and entropy effective rank were reported, and ridge regularization, eigenvalue flooring, shrinkage, pairwise clumping, leave-one-variant-out analyses and off-diagonal perturbations were evaluated as numerical robustness checks.",
        "Potential participant overlap was bounded analytically because the molecular-QTL donor set was small relative to the outcome GWASs. The maximal scenario allowed all molecular-QTL donors to overlap each outcome sample. Estimates were recomputed across the admissible overlap range and compared with the no-overlap estimate; the analysis quantified possible weak-instrument bias but could not identify unreported participant-level overlap."
    ]),
    ("Bidirectional MR, multivariable models and mediation identification", [
        "Genome-wide bidirectional MR evaluated lipid-to-AD and AD-to-lipid directions using independent genome-wide instruments and random-effects inverse-variance weighting. Heterogeneity, weighted-median, MR-Egger and extended-APOE sensitivity estimates were retained. A joint multivariable model evaluated the direct effects of low-density lipoprotein cholesterol, high-density lipoprotein cholesterol and triglycerides on AD with conditional instrument-strength diagnostics.",
        "At TSPAN14, local two-exposure multivariable models included exact exon5-6 usage and one cholesterol trait. Because the same correlated variants instrument both exposures, conditional F statistics were used as the primary identification gate. Principal-component generalized method-of-moments models decomposed the local linkage-disequilibrium matrix across retained dimensions. Direct splice and lipid estimates, lipid-to-splice estimates, indirect-effect estimates and attenuation fractions were examined across dimensions. A mediation result was classified as confirmatory only when all strength criteria and directional steps were satisfied. Independent trans-sQTL resources were also searched for an instrument mapping to the exact exon5-6 event."
    ]),
    ("Neural cell context and disease-state RNA", [
        "Cell localization, cell-type molecular QTLs and AD-associated RNA abundance were kept separate. The normalized canonical-branch sQTL was observed in an isolated microglial resource. The complementary donor-level count analysis arose from bulk anterior cingulate cortex and was not assigned to microglia, neurons or another constituent cell population. Coordinate-identical observations in additional bulk brain and spinal-cord tissues established neural-tissue consistency but were likewise not assigned to a specific cell type. Single-nucleus expression-QTL records were used only to show additional cellular contexts in which the TSPAN14 regulatory haplotype alters total expression.",
        "Disease-state analyses included a cross-study single-nucleus meta-analysis, adjusted SEA-AD pseudobulk models and independent microglial datasets. Study-level effects were retained on their reported scale and false-discovery correction was applied within the corresponding analysis. These tests evaluated whether AD status produces a reproducible change in total TSPAN14 RNA; they were not used as substitutes for genotype-dependent splice-QTL evidence."
    ]),
    ("Transcript, structure and protein-context interpretation", [
        "The exact and adjacent splice boundaries were mapped to the canonical TSPAN14-207 transcript and to the annotated TspanC8 large extracellular region. Residue-level predicted local distance difference test scores were extracted from the AlphaFold reference model to assess confidence in the location of amino acids 150/151. Structural confidence was interpreted as support for boundary localization, not as a prediction that the observed splice shift changes protein conformation.",
        "Public protein-QTL and proteomic resources were audited for isoform-resolved TSPAN14 or linked ADAM10 evidence. Published perturbation of the linked rs7922621 regulatory context was retained as an external functional anchor at the haplotype-to-gene level. Downstream effects on TSPAN14 isoform composition, ADAM10 trafficking and substrate-selective cleavage remain experimentally testable consequences of the exact splice association."
    ]),
    ("Statistical reporting and reproducibility", [
        "All tests were two-sided unless specified by the source analysis. Effect estimates are accompanied by standard errors, confidence intervals or posterior probabilities as appropriate. False-discovery rates are reported for transcriptomic analyses that tested multiple cell types or contrasts. Scripts consume machine-readable derived tables and generate the supplementary figures without manual modification of plotted values. A reviewer-ready archive contains the analysis scripts, derived source data, environment information, signed linkage-disequilibrium inputs and claim-to-code map; the repository address is reported in the manuscript Data and Code Availability statement."
    ]),
]

FIG_LEGENDS = [
    "Supplementary Figure S1 | APOE-conditioned sensitivity of genome-wide AD-lipid correlation. A, Baseline linkage disequilibrium score regression estimates for AD and five lipid traits; points show genetic correlations and horizontal bars show 95% confidence intervals. B, AD-HDL-C estimates after conditioning the extended APOE linkage-disequilibrium block on trait-specific or union lead signals and after an extreme 5-Mb physical-window exclusion. The window analysis is shown as sensitivity evidence rather than a formal conditional model.",
    "Supplementary Figure S2 | Complete non-APOE regional screen. A, Posterior-evidence matrix for all screened locus-trait pairs; point size and colour encode the posterior probability of a shared association (PP.H4). B, Ranked regional PP.H4 values, with the TSPAN14 locus distinguished from other screened regions and the 0.80 reference threshold shown by the vertical dashed line.",
    "Supplementary Figure S3 | Colocalization and fine-mapping diagnostics at TSPAN14. A, Trait-specific posterior inclusion probabilities for the 15 highest-ranked variants. B, Sizes and top posterior inclusion probabilities of SuSiE credible sets. C, Default-prior colocalization posterior probabilities for the exact exon5-6 event with AD and cholesterol traits; the axis is restricted to the high-posterior range to resolve differences among models. Conservative-prior sensitivity is reported in Supplementary Table S4.",
    "Supplementary Figure S4 | Genotype-dependent canonical–cryptic acceptor balance. A, Donor-level canonical read fraction within the local two-acceptor cluster by rs7080009 AD-risk C-allele dosage; points represent donors and horizontal bars show medians. B, Odds ratios per risk allele from Firth donor-detection models and the beta-binomial read-count model; horizontal bars show profile-likelihood 95% confidence intervals. C, Firth odds ratios after progressively increasing the minimum local-cluster read depth; point size denotes the retained donor count.",
    "Supplementary Figure S5 | Coordinate-identical exon5-6 sQTL consistency across partially overlapping neural tissues. A, Signed -log10(P) values for six risk-aligned variants across anterior cingulate cortex, hippocampus, putamen and cervical spinal cord; positive values denote the same direction as the primary event. B, Median risk-aligned normalized effect size by tissue, with point size proportional to the strongest association evidence. GTEx donor overlap precludes interpretation as four independent replications.",
    "Supplementary Figure S6 | Relationship between adjacent TSPAN14 junctions. A, Spearman correlations between exon5-6 and exon6-7 junction counts by neural tissue; point size denotes the number of samples in which both junctions were observed. B, Numbers of samples, donors and jointly observed junction pairs contributing to each tissue estimate. Correlation of aggregate counts does not establish co-occurrence within a single RNA molecule.",
    "Supplementary Figure S7 | Rebuilt-LD exact-event cis-MR robustness. A, Risk-aligned local estimates using the signed European-ancestry linkage-disequilibrium matrix; horizontal bars show 95% confidence intervals. B, Median estimates and 2.5th-97.5th percentile ranges across 1,000 off-diagonal linkage-disequilibrium perturbations at each noise level. C, Eigenvalue, condition-number and effective-rank diagnostics for the local correlation matrix. These analyses evaluate directional and numerical stability rather than independent causal identification.",
    "Supplementary Figure S8 | Causal-scope analyses. A, Genome-wide bidirectional MR estimates for AD and lipid traits; red points denote nominal P<0.05 and grey points denote P>=0.05. B, Estimated indirect effects across retained linkage-disequilibrium principal components in local PC-GMM models. C, Minimum instrument-strength statistic across each model dimension; the dashed line marks F=10. Instability across dimensions and weak conditional strength prevent confirmatory mediation inference.",
    "Supplementary Figure S9 | Cell context, disease-state RNA and structural localization. A, Evidence matrix separating exact splice QTL, total-expression QTL, disease-state and perturbation observations across neural contexts. B, Selected disease-state estimates from single-nucleus meta-analysis and adjusted SEA-AD pseudobulk analyses; intervals are shown where available, and none passed a 5% false-discovery-rate threshold. C, AlphaFold reference-model confidence across the TSPAN14 large extracellular region, with the exact amino-acid 150/151 boundary marked."
]


def set_cell_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "6"), ("space", "1"), ("color", "17365D")):
        bottom.set(qn(f"w:{key}"), value)
    borders.append(bottom); ppr.append(borders)


def add_text(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p


def build():
    doc = Document(); sec = doc.sections[0]
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.78); sec.right_margin=Inches(.78)
    styles=doc.styles
    styles["Normal"].font.name="Times New Roman"; styles["Normal"].font.size=Pt(10)
    for name,size,color in (("Heading 1",15,"17365D"),("Heading 2",11,"17365D")):
        styles[name].font.name="Times New Roman"; styles[name].font.size=Pt(size); styles[name].font.bold=True; styles[name].font.color.rgb=RGBColor.from_string(color)
    cover=doc.add_paragraph(); cover.alignment=WD_ALIGN_PARAGRAPH.CENTER; cover.paragraph_format.space_before=Pt(70)
    run=cover.add_run(TITLE); run.font.name="Times New Roman"; run.font.size=Pt(19); run.bold=True; run.font.color.rgb=RGBColor.from_string("17365D")
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_before=Pt(18)
    r=sub.add_run("Supplementary Information"); r.font.name="Times New Roman"; r.font.size=Pt(15); r.bold=True
    auth=doc.add_paragraph("Xinpeng Yang and Xindong Shui"); auth.alignment=WD_ALIGN_PARAGRAPH.CENTER
    summary=doc.add_paragraph("Supplementary Methods | Supplementary Figures S1-S9 | Supplementary Tables S1-S19"); summary.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    h=doc.add_heading("Supplementary Methods",level=1); set_cell_border(h)
    for heading, paragraphs in METHODS:
        doc.add_heading(heading,level=2)
        for text in paragraphs: add_text(doc,text)
    doc.add_page_break()
    h=doc.add_heading("Supplementary Figures",level=1); set_cell_border(h)
    for i, legend in enumerate(FIG_LEGENDS,1):
        if i>1: doc.add_page_break()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(6)
        p.add_run().add_picture(str(BASE/"figures"/f"Supplementary_Figure_S{i}.png"),width=Inches(6.65))
        cap=doc.add_paragraph(); cap.paragraph_format.keep_together=True; cap.paragraph_format.space_before=Pt(4); cap.paragraph_format.space_after=Pt(0)
        prefix=f"Supplementary Figure S{i} |"; cap.add_run(prefix).bold=True; cap.add_run(legend.split("|",1)[1])
    doc.add_page_break()
    h=doc.add_heading("Supplementary Table Legends",level=1); set_cell_border(h)
    index=pd.read_csv(BASE/"source_tables"/"Table_S00_Index.tsv",sep="\t")
    add_text(doc,"Supplementary Tables S1-S19 are supplied in the accompanying Excel workbook. Each worksheet was rebuilt from the current audited analysis outputs; blank cells indicate non-applicable fields and numerical values retain source precision.")
    for row in index.itertuples(index=False):
        add_text(doc,f"Supplementary Table S{row.number}. {row.title}.",bold_prefix=f"Supplementary Table S{row.number}.")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Supplementary Information | ")
    field=OxmlElement("w:fldSimple"); field.set(qn("w:instr"),"PAGE"); footer._p.append(field)
    BASE.mkdir(parents=True,exist_ok=True)
    doc.save(BASE/"Supplementary_Information.docx")
    print(BASE/"Supplementary_Information.docx")


if __name__=="__main__":
    import pandas as pd
    build()
