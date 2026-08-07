#!/usr/bin/env python3
"""Verify the synchronized S1-S19 supplementary release."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "outputs" / "supplement_v19"
TITLE = "A non-APOE Alzheimer disease-cholesterol locus converges on TSPAN14 splice choice"


def document_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def referenced_numbers(body: str, kind: str) -> set[int]:
    label = r"Fig(?:ure)?s?" if kind == "figure" else r"Tables?"
    pattern = re.compile(rf"Supplementary\s+{label}\.?\s+([^.;\)]+)", re.I)
    numbers: set[int] = set()
    for match in pattern.finditer(body):
        block = match.group(1)
        for start, end in re.findall(r"S?(\d+)\s*[-–]\s*S?(\d+)", block, re.I):
            numbers.update(range(int(start), int(end) + 1))
        numbers.update(int(value) for value in re.findall(r"S(\d+)", block, re.I))
    return numbers


def office_zip_ok(path: Path) -> bool:
    with ZipFile(path) as archive:
        return archive.testzip() is None


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manuscript", type=Path, required=True)
    args = parser.parse_args()

    expected_figures = set(range(1, 10))
    expected_tables = set(range(1, 20))
    manuscript = document_text(args.manuscript)
    supplement_docx = BASE / "Supplementary_Information.docx"
    supplement_pdf = BASE / "Supplementary_Information.pdf"
    workbook = BASE / "Supplementary_Tables.xlsx"
    supplement = document_text(supplement_docx)

    index = pd.read_csv(BASE / "source_tables" / "Table_S00_Index.tsv", sep="\t")
    table_shapes = {}
    shape_errors = []
    for row in index.itertuples(index=False):
        table_path = BASE / "source_tables" / row.file
        data = pd.read_csv(table_path, sep="\t")
        table_shapes[f"S{row.number}"] = [len(data), len(data.columns)]
        if len(data) != int(row.rows) or len(data.columns) != int(row.columns):
            shape_errors.append(f"S{row.number}: index={row.rows}x{row.columns}, file={len(data)}x{len(data.columns)}")

    s8 = pd.read_csv(BASE / "source_tables" / "Table_S08.tsv", sep="\t")
    s8_blocks = s8["analysis_block"].astype(str).str.strip()
    s8_ok = (
        len(s8) == 38
        and (s8_blocks == "GTEx focal-tissue donor counts").sum() == 4
        and (s8_blocks == "GTEx focal-tissue donor overlap").sum() == 6
        and not s8.astype(str).apply(lambda column: column.str.contains("risk-aligned replication", case=False, regex=False)).any().any()
    )

    required = [supplement_docx, supplement_pdf, workbook, BASE / "source_tables" / "Table_S00_Index.tsv"]
    required += [BASE / "source_tables" / f"Table_S{i:02d}.tsv" for i in expected_tables]
    required += [BASE / "figures" / f"Supplementary_Figure_S{i}.{ext}" for i in expected_figures for ext in ("png", "pdf", "svg", "tiff")]

    manuscript_figures = referenced_numbers(manuscript, "figure")
    manuscript_tables = referenced_numbers(manuscript, "table")
    figure_legends = {int(value) for value in re.findall(r"Supplementary Figure S(\d+)", supplement)}
    table_legends = {int(value) for value in re.findall(r"Supplementary Table S(\d+)", supplement)}

    report = {
        "title_exact": TITLE in supplement.splitlines()[:5],
        "manuscript_figure_references": sorted(manuscript_figures),
        "manuscript_table_references": sorted(manuscript_tables),
        "figure_legends": sorted(figure_legends),
        "table_legends": sorted(table_legends),
        "missing_or_empty_files": [str(path) for path in required if not path.exists() or path.stat().st_size == 0],
        "office_zip_integrity": {supplement_docx.name: office_zip_ok(supplement_docx), workbook.name: office_zip_ok(workbook)},
        "pdf_pages": len(PdfReader(supplement_pdf).pages),
        "workbook_qa": json.loads((BASE / "qa" / "workbook" / "workbook_qa.json").read_text(encoding="utf-8"))["status"],
        "source_table_shapes": table_shapes,
        "source_table_shape_errors": shape_errors,
        "s8_deduplicated_and_relabelled": s8_ok,
        "stale_terms": sorted(set(re.findall(r"(?:independent exact-event bulk-tissue replication|risk-aligned replication|replication across four neural tissues)", supplement + "\n" + "\n".join(path.read_text(encoding="utf-8") for path in (BASE / "source_tables").glob("*.tsv")), re.I))),
    }
    report["missing_manuscript_figures"] = sorted(expected_figures - manuscript_figures)
    report["missing_manuscript_tables"] = sorted(expected_tables - manuscript_tables)
    report["missing_figure_legends"] = sorted(expected_figures - figure_legends)
    report["missing_table_legends"] = sorted(expected_tables - table_legends)
    failures = [
        not report["title_exact"], report["missing_or_empty_files"],
        not all(report["office_zip_integrity"].values()), report["pdf_pages"] != 14,
        report["workbook_qa"] != "PASS", shape_errors, not s8_ok, report["stale_terms"],
        report["missing_manuscript_figures"], report["missing_manuscript_tables"],
        report["missing_figure_legends"], report["missing_table_legends"],
    ]
    report["status"] = "FAIL" if any(failures) else "PASS"
    (BASE / "qa_report.json").write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(report, indent=2, ensure_ascii=False))
    raise SystemExit(0 if report["status"] == "PASS" else 1)


if __name__ == "__main__":
    main()
