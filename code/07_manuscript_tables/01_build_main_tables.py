#!/usr/bin/env python3
"""Rebuild the two main manuscript tables as compact SCI three-line tables."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TABLES = [
    {
        "title": "Table 1 | APOE-aware genetic screening and prioritization of the TSPAN14 locus",
        "headers": [
            "Stage",
            "Trait pair",
            "Primary result",
            "Robustness analysis",
            "Biological interpretation",
        ],
        "widths": [1.25, 1.15, 1.45, 2.55, 3.90],
        "center_cols": {1, 2},
        "rows": [
            [
                "Genome-wide screening",
                "AD–HDL-C",
                "Genetic correlation = 0.1394 (SE = 0.0436); P = 0.0014",
                "Extended-APOE LD conditioning: correlation = 0.1155 (own lead) and 0.1181 (pair union), both P = 0.0002; ±5-Mb exclusion: correlation = 0.1059, P = 0.0001",
                "The AD–HDL-C correlation is sensitive to, but not fully explained by, the extended APOE association block.",
            ],
            ["Genome-wide screening", "AD–LDL-C", "Genetic correlation = −0.0227 (SE = 0.0310); P = 0.4629", "No significant genome-wide correlation", "Not prioritized by the global screen."],
            ["Genome-wide screening", "AD–TG", "Genetic correlation = −0.0351 (SE = 0.0320); P = 0.2739", "No significant genome-wide correlation", "Not prioritized by the global screen."],
            ["Genome-wide screening", "AD–TC", "Genetic correlation = −0.0054 (SE = 0.0344); P = 0.8741", "No significant genome-wide correlation", "Not prioritized by the global screen."],
            ["Genome-wide screening", "AD–non-HDL-C", "Genetic correlation = −0.0638 (SE = 0.0399); P = 0.1093", "No significant genome-wide correlation", "Not prioritized by the global screen."],
            [
                "TSPAN14 locus",
                "AD–TC",
                "Regional PP.H4 = 0.9757",
                "coloc.susie PP.H4 = 0.9587; HyPrColoc = 0.9413; exact exon5–6 sQTL PP.H4 = 0.9771",
                "Most prior-robust local convergence, with the exact TSPAN14 splice feature as a candidate molecular readout.",
            ],
            [
                "TSPAN14 locus",
                "AD–LDL-C",
                "Regional PP.H4 = 0.9619",
                "coloc.susie PP.H4 = 0.9507; HyPrColoc = 0.9099; exact exon5–6 sQTL PP.H4 = 0.9589",
                "High default-prior support, but sensitive to a tenfold lower shared-signal prior.",
            ],
            [
                "TSPAN14 locus",
                "AD–non-HDL-C",
                "Regional PP.H4 = 0.9586",
                "coloc.susie PP.H4 = 0.9571; HyPrColoc = 0.9025; exact exon5–6 sQTL PP.H4 = 0.9598",
                "High default-prior support, but sensitive to a tenfold lower shared-signal prior.",
            ],
        ],
        "note": (
            "Note. AD, Alzheimer disease; APOE, apolipoprotein E; HDL-C, high-density lipoprotein cholesterol; "
            "LD, linkage disequilibrium; LDL-C, low-density lipoprotein cholesterol; PP.H4, posterior probability "
            "of a shared association hypothesis; SE, standard error; sQTL, splicing "
            "quantitative trait locus; TC, total cholesterol; TG, triglycerides. LD conditioning and physical-window "
            "exclusion are complementary sensitivity analyses and do not constitute direct adjustment for APOE genotype."
        ),
    },
    {
        "title": "Table 2 | Replication, effect magnitude and molecular triangulation for exact TSPAN14 exon5–6 splicing",
        "headers": ["Evidence", "Tissue or outcome", "Effect estimate", "Statistical support", "Interpretation"],
        "widths": [1.55, 1.45, 1.80, 2.20, 3.30],
        "center_cols": {1, 2},
        "rows": [
            ["Coordinate-matched canonical branch", "Anterior cingulate cortex (BA24)", "Risk-aligned NES = 1.159–1.219", "248 records; minimum P = 3.94 × 10⁻²¹; 227 at P ≤ 5 × 10⁻⁸", "Coordinate-matched canonical-branch association; tissue records are not independent variants or cohorts."],
            ["Coordinate-matched canonical branch", "Hippocampus", "Risk-aligned NES = 0.636–0.782", "30 records; minimum P = 2.61 × 10⁻¹¹; 18 at P ≤ 5 × 10⁻⁸", "Cross-tissue consistency in GTEx; donor overlap is reported separately."],
            ["Coordinate-matched canonical branch", "Putamen", "Risk-aligned NES = 0.538–0.705", "32 records; minimum P = 3.54 × 10⁻¹²; 25 at P ≤ 5 × 10⁻⁸", "Cross-tissue consistency in GTEx; donor overlap is reported separately."],
            ["Coordinate-matched canonical branch", "Cervical spinal cord", "Risk-aligned NES = 0.828–0.900", "34 records; minimum P = 2.24 × 10⁻¹²; 24 at P ≤ 5 × 10⁻⁸", "Cross-tissue consistency in GTEx; donor overlap is reported separately."],
            ["Canonical–cryptic read balance", "BA24 raw junction counts (n = 147; rs7080009 C-risk allele)", "Canonical fraction: 95.57% in T/T and 100% in C/C donors; descriptive difference = 4.43 percentage points", "Beta-binomial OR = 0.1499 (95% CI, 0.0918–0.2375; P = 4.45 × 10⁻¹⁴); depth-adjusted Firth OR = 0.0131 (0.00141–0.0539; P = 1.04 × 10⁻¹⁷)", "Count and donor-detection models support genotype-dependent acceptor balance; the fraction is not ΔPSI or a full-length isoform proportion."],
            ["Canonical-branch colocalization", "AD", "Default-prior PP.H4 = 0.9909", "coloc.susie PP.H4 = 0.9811; p12/10 PP.H4 = 0.9155", "The canonical-branch association pattern is compatible with the local AD configuration and remains prior robust."],
            ["Canonical-branch colocalization", "TC; LDL-C; non-HDL-C", "Default-prior PP.H4 = 0.9771; 0.9589; 0.9598", "p12/10 PP.H4 = 0.8098; 0.6999; 0.7050", "TC remains prior robust; LDL-C and non-HDL-C are prior sensitive."],
            ["LD-aware cis-MR", "AD", "β = 0.0651 (SE = 0.0108)", "P = 1.78 × 10⁻⁹; sentinel Wald β = 0.0653 (SE = 0.0110)", "Positive direction across LD sensitivity analyses; locus-level evidence, not causal proof."],
            ["LD-aware cis-MR", "TC", "β = 0.00881 (SE = 0.00165)", "P = 9.65 × 10⁻⁸", "Positive direction across LD sensitivity analyses; locus-level evidence."],
            ["LD-aware cis-MR", "LDL-C", "β = 0.00833 (SE = 0.00167)", "P = 5.77 × 10⁻⁷", "Positive direction across LD sensitivity analyses; locus-level evidence."],
            ["LD-aware cis-MR", "Non-HDL-C", "β = 0.00972 (SE = 0.00196)", "P = 7.44 × 10⁻⁷; prior Q P = 0.0041; Egger-intercept P = 0.0027", "Positive estimate with heterogeneity and directional-pleiotropy sensitivity flags; downgraded evidence."],
            ["Overlap sensitivity", "Complete-overlap upper-bound scenario for 147 sQTL donors", "Rebuilt signed-LD baseline retained; no overlap-adjusted point estimate inferred", "Minimum instrument F = 32.49; relative weak-instrument bias bound = 3.08%", "Conservative bound only because cohort-resolved sampling covariance was unavailable."],
        ],
        "note": (
            "Note. AD, Alzheimer disease; BA24, Brodmann area 24; CI, confidence interval; HC3, heteroscedasticity-consistent "
            "standard error; IER, intron excision ratio; LD, linkage disequilibrium; LDL-C, low-density lipoprotein cholesterol; "
            "MR, Mendelian randomization; NES, normalized effect size; PP.H4, posterior probability of a shared association "
            "hypothesis; SE, standard error; sQTL, splicing quantitative trait locus; TC, total cholesterol. The canonical "
            "read fraction is calculated within a local two-acceptor cluster and is not conventional exon-inclusion ΔPSI. "
            "Default-prior PP.H4 values and conservative-prior sensitivity are shown; complete details are reported in "
            "Supplementary Table S4. Correlated cis variants provide locus-level sensitivity evidence rather than independent "
            "experiments or proof of circulating-lipid mediation."
        ),
    },
    {
        "title": "Table 3 | Causal-scope analyses of the TSPAN14–lipid–AD association",
        "headers": ["Question", "Design", "Identification strength", "Principal result", "Inference"],
        "widths": [1.55, 2.15, 1.85, 2.20, 2.55],
        "center_cols": set(),
        "rows": [
            ["Do circulating lipids have a genome-wide total effect on AD?", "Lipid → AD arm of bidirectional two-sample MR", "700 TC, 593 LDL-C, 547 non-HDL-C, 805 HDL-C and 692 TG instruments", "TC, LDL-C and non-HDL-C: P = 0.535, 0.717 and 0.905; HDL-C: P = 4.09 × 10⁻⁶; TG: P = 0.049; the latter signals showed substantial heterogeneity", "No robust general lipid-to-AD pathway was resolved across traits and sensitivity analyses."],
            ["Do lipid fractions have mutually adjusted direct effects on AD?", "Genome-wide MVMR of LDL-C, HDL-C and TG", "1,327 instruments; minimum conditional F = 20.86", "All direct-effect P ≥ 0.152", "No independently identified direct lipid effect after mutual adjustment."],
            ["Can local lipid and splice effects be separated?", "Strict and relaxed locus-restricted MVMR", "Strict model: one instrument, underidentified; relaxed model: six instruments, conditional F = 2.18–3.71", "All relaxed-model P ≥ 0.101", "Dense local LD prevents stable separation of lipid and splice coefficients."],
            ["Is TC → exon5–6 splicing → AD compatible with the data?", "Three-PC locus-restricted PC-GMM", "Three PCs captured the estimable local LD dimensions", "TC coefficient attenuated by 69.3%; lipid-to-splice P = 0.069; indirect-effect P = 0.110", "Compatible with, but does not identify, serial mediation."],
            ["Is the mediation result stable to model dimension?", "PC-GMM sweep across 3–15, 20, 25 and 30 PCs", "Prespecified strength, significance and stability criteria", "0 models passed all confirmatory criteria", "Attenuation is dimension-sensitive and not confirmatory."],
            ["Can independent two-step mediation be tested?", "Audit of INTERVAL, BrainMeta and GTEx trans-sQTL resources", "No instrument for the exact exon5–6 event", "Two-step MR not identifiable", "Current public data distinguish local convergence from proven lipid-to-splice-to-AD mediation."],
        ],
        "note": (
            "Note. AD, Alzheimer disease; HDL-C, high-density lipoprotein cholesterol; IVW, inverse-variance weighted; "
            "LD, linkage disequilibrium; LDL-C, low-density lipoprotein cholesterol; MR, Mendelian randomization; MVMR, "
            "multivariable Mendelian randomization; PC-GMM, principal-component generalized method of moments; TC, total "
            "cholesterol; TG, triglycerides. Absence from significance-filtered trans-sQTL resources was treated as censoring, "
            "not as a zero association."
        ),
    },
    {
        "title": "Table 4 | Cell-context and functional interpretation of TSPAN14 regulation",
        "headers": ["Evidence layer", "Source or context", "Principal result", "Role in the mechanism"],
        "widths": [1.65, 2.40, 3.35, 2.90],
        "center_cols": set(),
        "rows": [
            ["Exact-event cellular localization", "MiGA isolated human microglia", "Six risk-aligned variants increased use of the exact exon5–6 junction (β = 0.989–1.251).", "Directly localizes the core splice association to microglia without inferring its source from bulk tissue."],
            ["Exact-event neural replication", "GTEx BA24, hippocampus, putamen and cervical spinal cord", "The identical exon5–6 junction showed concordant risk-allele effects in all four tissues.", "Establishes reproducibility across neural tissues while leaving bulk-tissue cellular origin unspecified."],
            ["Single-nucleus regulatory context", "NIAGADS ROSMAP/CUIMC/MIT eQTL analyses", "Three core risk variants increased TSPAN14 expression in astrocytic and excitatory-neuronal contexts (β = 0.233–0.378; FDR ≤ 1.13 × 10⁻⁴).", "Shows that the regulatory haplotype is active in multiple brain cell contexts; it is not an AD-state differential-expression result."],
            ["AD-state total RNA", "Single-cell meta-analysis, SEA-AD and independent pseudobulk datasets", "No reproducible FDR-significant change in total TSPAN14 expression across microglia or neuronal populations.", "Favors genotype-linked transcript processing over a universal late-stage change in total gene abundance."],
            ["Genotype-resolved perturbation", "Published rs7922621 prime editing and enhancer CRISPRi", "rs7922621 is in high LD with rs7080009 (r² = 0.987); perturbation altered TSPAN14, cell-surface ADAM10 and soluble TREM2.", "Provides an external functional anchor for the regulatory block and the TSPAN14–ADAM10–TREM2 axis."],
            ["Protein-region mapping", "TSPAN14 canonical transcript, UniProt and AlphaFold reference model", "The exon5–6 boundary maps to AA150/151 within EC2 (AA114–232); local reference pLDDT = 93.3.", "Defines a structurally resolved interface hypothesis for isoform-specific functional testing."],
        ],
        "note": (
            "Note. AD, Alzheimer disease; ADAM10, a disintegrin and metalloproteinase domain-containing protein 10; "
            "BA24, Brodmann area 24; EC2, large extracellular loop; eQTL, expression quantitative trait locus; FDR, false "
            "discovery rate; LD, linkage disequilibrium; MiGA, Microglia Genomic Atlas; pLDDT, predicted local distance "
            "difference test; TREM2, triggering receptor expressed on myeloid cells 2. Exact splice-to-protein consequences "
            "remain experimentally testable predictions."
        ),
    },
]

# The current journal-format manuscript retains Tables 1 and 2 in the main text;
# causal-scope and cell-context detail are reported in the supplementary tables.
TABLES = TABLES[:2]


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = "w:" + edge_name
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches):
    widths_dxa = [round(x * 1440) for x in widths_inches]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_inches[idx])


def style_run(run, size=8.0, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = None


def style_paragraph(paragraph, *, size=8.0, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold, italic=italic)


def format_rg_subscript(paragraph, *, size=8.0):
    """Render r_g as an italic r with a true subscript g in table cells."""
    text = paragraph.text
    if "r_g" not in text:
        return
    alignment = paragraph.alignment
    paragraph.clear()
    parts = text.split("r_g")
    for idx, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            style_run(run, size=size)
        if idx < len(parts) - 1:
            run = paragraph.add_run("r")
            style_run(run, size=size, italic=True)
            run = paragraph.add_run("g")
            style_run(run, size=size, italic=True)
            run.font.subscript = True
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def style_table(table, center_cols):
    none = {"val": "nil"}
    top_bottom = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
    header_rule = {"val": "single", "sz": "8", "color": "000000", "space": "0"}

    # Remove all borders before adding the three required horizontal rules.
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, start=none, end=none, insideH=none, insideV=none)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)

    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top_bottom, bottom=header_rule)
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in table.rows[1:]:
        for col_idx, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, size=8.0, align=align)

    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=top_bottom)


def replace_table(doc, old_table, spec):
    new_table = doc.add_table(rows=1, cols=len(spec["headers"]))
    new_table.style = "Normal Table"
    for idx, value in enumerate(spec["headers"]):
        new_table.rows[0].cells[idx].text = value
    for values in spec["rows"]:
        cells = new_table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    set_table_geometry(new_table, spec["widths"])
    style_table(new_table, spec["center_cols"])
    old_table._tbl.addprevious(new_table._tbl)
    old_table._element.getparent().remove(old_table._element)
    return new_table


def find_paragraph(doc, prefix):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}, found {len(matches)}")
    return matches[0]


def format_caption(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    style_run(run, size=9.0, bold=True)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True


def format_note(paragraph, text):
    paragraph.clear()
    lead, rest = text.split(" ", 1)
    run = paragraph.add_run(lead + " ")
    style_run(run, size=8.0, italic=True)
    run = paragraph.add_run(rest)
    style_run(run, size=8.0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0


def write_source_tables(out_dir):
    out_dir.mkdir(parents=True, exist_ok=True)
    for idx, spec in enumerate(TABLES, start=1):
        path = out_dir / f"Table_{idx}_source.tsv"
        with path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle, delimiter="\t")
            writer.writerow(spec["headers"])
            writer.writerows(spec["rows"])


def citation_audit(doc):
    first_positions = {}
    all_hits = {str(i): [] for i in range(1, 3)}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if text.startswith("Table ") or text.startswith("Note."):
            continue
        for table_no in range(1, 3):
            token = f"Table {table_no}"
            if token in text:
                all_hits[str(table_no)].append(idx)
                first_positions.setdefault(table_no, idx)
    if sorted(first_positions) != [1, 2]:
        raise RuntimeError(f"Missing in-text table citation: {first_positions}")
    order = [first_positions[i] for i in range(1, 3)]
    if order != sorted(order):
        raise RuntimeError(f"In-text table citations are not sequential: {first_positions}")
    return {"first_paragraph": first_positions, "all_paragraphs": all_hits}


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--audit-dir", required=True, type=Path)
    args = parser.parse_args()

    doc = Document(args.source)
    if len(doc.tables) != 2:
        raise RuntimeError(f"Expected two source tables, found {len(doc.tables)}")
    expected_headers = ["Stage", "Evidence"]
    actual_headers = [table.rows[0].cells[0].text.strip() for table in doc.tables]
    if actual_headers != expected_headers:
        raise RuntimeError(f"Unexpected source table headers: {actual_headers}")

    old_tables = list(doc.tables)
    for idx, (old_table, spec) in enumerate(zip(old_tables, TABLES), start=1):
        format_caption(find_paragraph(doc, f"Table {idx} |"), spec["title"])
        # Each table is immediately followed by its note in the manuscript table section.
        note_node = old_table._tbl.getnext()
        while note_node is not None and note_node.tag != qn("w:p"):
            note_node = note_node.getnext()
        note_paragraph = next((p for p in doc.paragraphs if p._p is note_node), None)
        if note_paragraph is None or not note_paragraph.text.strip().startswith("Note."):
            raise RuntimeError(f"Could not identify note following Table {idx}")
        format_note(note_paragraph, spec["note"])
        replace_table(doc, old_table, spec)

    citations = citation_audit(doc)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    write_source_tables(args.audit_dir / "source_data")

    audit = {
        "source": str(args.source.resolve()),
        "output": str(args.output.resolve()),
        "table_count": len(doc.tables),
        "table_dimensions": [
            {"table": idx, "rows": len(table.rows), "columns": len(table.columns)}
            for idx, table in enumerate(doc.tables, start=1)
        ],
        "citation_audit": citations,
        "titles": [spec["title"] for spec in TABLES],
    }
    args.audit_dir.mkdir(parents=True, exist_ok=True)
    (args.audit_dir / "main_table_rebuild_audit.json").write_text(
        json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
