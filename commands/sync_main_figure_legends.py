from pathlib import Path

from docx import Document
from docx.shared import Pt


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 2.0


def main() -> None:
    package_root = Path(__file__).resolve().parents[1]
    submission_root = package_root.parent
    source = submission_root / "ADJ-D-26-02420_Revised_Manuscript_Round9_TNR_Superscript_Citations.docx"
    destination = submission_root / "ADJ-D-26-02420_Revised_Manuscript_Round10_Figures_Synced.docx"

    document = Document(source)
    legends = {
        number: (package_root / "figures" / f"Figure_{number:02d}" / "legend.md").read_text(encoding="utf-8").strip()
        for number in range(1, 7)
    }

    replaced = set()
    for paragraph in document.paragraphs:
        for number, legend in legends.items():
            if paragraph.text.startswith(f"Figure {number} |"):
                replace_paragraph_text(paragraph, legend)
                replaced.add(number)
                break

    if replaced != set(legends):
        missing = sorted(set(legends) - replaced)
        raise RuntimeError(f"Missing figure legends in manuscript: {missing}")

    document.save(destination)
    print(destination)


if __name__ == "__main__":
    main()
