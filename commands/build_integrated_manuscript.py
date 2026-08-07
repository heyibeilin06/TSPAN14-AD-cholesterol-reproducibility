#!/usr/bin/env python3
"""Insert canonical main figures into a reviewer-readable manuscript copy."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--figure-root", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--height-inches", type=float, default=5.95)
    args = parser.parse_args()

    document = Document(args.input)
    legends = {
        int(match.group(1)): paragraph
        for paragraph in document.paragraphs
        if (match := re.match(r"^Figure ([1-6]) \|", paragraph.text))
    }
    if sorted(legends) != list(range(1, 7)):
        raise SystemExit(f"Expected Figure 1-6 legends; found {sorted(legends)}")

    headings = [p for p in document.paragraphs if p.text.strip() == "Figure legends"]
    if len(headings) != 1:
        raise SystemExit(f"Expected one Figure legends heading; found {len(headings)}")
    headings[0].paragraph_format.page_break_before = True
    headings[0].paragraph_format.keep_with_next = True

    for number in range(1, 7):
        image_path = (
            args.figure_root
            / f"Figure_{number:02d}/output/Figure_{number}.png"
        )
        if not image_path.is_file():
            raise SystemExit(f"Missing canonical figure image: {image_path}")

        legend = legends[number]
        figure_paragraph = legend.insert_paragraph_before()
        figure_paragraph.paragraph_format.page_break_before = number != 1
        figure_paragraph.paragraph_format.space_before = 0
        figure_paragraph.paragraph_format.space_after = 0
        figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = figure_paragraph.add_run().add_picture(
            str(image_path), height=Inches(args.height_inches)
        )
        set_alt_text(
            shape,
            f"Figure {number}",
            f"Main manuscript Figure {number}; full interpretation is provided in the following legend.",
        )

        # Keep the figure at readable size and place its double-spaced legend
        # on the following page instead of leaving a one-line orphan.
        legend.paragraph_format.page_break_before = True
        legend.paragraph_format.keep_together = True

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
